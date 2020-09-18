from django.shortcuts import render, HttpResponse,redirect, get_object_or_404
import requests
from requests  import get
from bs4 import BeautifulSoup
from django.http import HttpResponseRedirect
from .models import Claimset
from selenium import webdriver 
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver import Chrome
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import os
import time
import io
from pathlib import Path
from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager

from .forms import ClaimForm
chrome_options = webdriver.ChromeOptions()
chrome_options.add_argument('--no-sandbox')
chrome_options.add_argument("--disable-extensions")
chrome_options.add_argument("--incognito")
#browser = webdriver.Chrome(executable_path=ChromeDriverManager().install(), chrome_options=chrome_options)

# Build paths inside the project like this: BASE_DIR / 'subdir'.
BASE_DIR = Path(__file__).resolve(strict=True).parent.parent

# PROJECT_ROOT = os.path.dirname(os.path.realpath('__file__'))
# print(PROJECT_ROOT)
# DRIVER_BIN = os.path.join(PROJECT_ROOT, "chromedriver")
# print(DRIVER_BIN)
chrome_options = webdriver.ChromeOptions()
chrome_options.add_argument( '--headless')
chrome_options.add_experimental_option("excludeSwitches", ['enable-automation'])
#driver = webdriver.Chrome(DRIVER_BIN, options=chrome_options) 
MEDIA_ROOT = os.path.join(BASE_DIR, 'media')

# def new(request):
#     form = ClaimForm()
#     context = {"form":form}
#     return render(request,"patoa/new.html", context)

# global AB112, searchterm,purl, cln
# def scrape_get(request):
#         #if request.method == "POST":
#             #filled_form =ClaimForm(request.POST)
#             form =ClaimForm()
#             # print(filled_form)
#             # #filled_form (initial = {"patno": "patno", "pubno": "pubno", "add112" :"ad112","obj" : "obj"})
#             # if filled_form.is_valid():
#             #     claims = Claimset()
#             #     claims.patno = filled_form.cleaned_data['patno']
#             #     claims.pubno = filled_form.cleaned_data['pubno']
#             #     claims.add112 = filled_form.cleaned_data['add112']
#             #     claims.obj = filled_form.cleaned_data['obj']
#             #     if claims.is_valid():
#             #         claims.save()
#             # if form.is_valid():
#             #     form.save()  
#             return redirect('/')
#             #return render(request, 'patoa/result.html', { 'form':form})
searchterm= [""]
def scrape_post(request):
    form = ClaimForm()
    if request.method == 'POST':
        filled_form = ClaimForm(request.POST)
        if filled_form.is_valid():
        #form.save()
        #Claimset = form.save(commit=False)
            claims = Claimset()
            claims.patno = filled_form.cleaned_data['patno']
            claims.pubno = filled_form.cleaned_data['pubno']
            claims.add112 = filled_form.cleaned_data['add112']
            claims.obj = filled_form.cleaned_data['obj']
            print("Pat appl no :",claims.patno)
            claims.save()

    # if form.is_valid():
    #     #form.save()
    #     #Claimset = form.save(commit=False)
    #     patno2 = form.cleaned_data.get("patno")
    #     pubno2 = form.cleaned_data.get("pubno")
    #     ADD112 = form.cleaned_data.get("add112")
    #     obj2 = form.cleaned_data.get("obj")
    #     print("Pat appl no :",patno2)
        if claims.patno:
            searchterm[0] = (claims.patno)
        else:
            searchterm[0] = claims.pubno
        print("You selected searchterm :", searchterm[0]) 

        if claims.add112 and claims.obj:
            print("You selected both")
            AB112 = '112obj'
        elif claims.add112:
            print("You selected 112")
            AB112 = '112'
        elif claims.obj :
            print("You selected obj")
            AB112= 'obj2'
        else:
            print("No selection")
            AB112 =''
        # if patno2:
        #     searchterm = patno2
        # else:
        #     searchterm = pubno2
        # print("You selected searchterm :",searchterm) 

        # if ADD112 and obj2:
        #     print("You selected both")
        #     AB112 = '112obj'
        # elif ADD112:
        #     print("You selected 112")
        #     AB112 = '112'
        # elif obj2 :
        #     print("You selected obj")
        #     AB112= 'obj2'
        # else:
        #     print("No selection")
        #     AB112 =''
        # print("AB112 is:\n", AB112)
            #print("path for driver :::", DRIVER_BIN)
            #driver = webdriver.Chrome(DRIVER_BIN, options=chrome_options)
        driver = webdriver.Chrome(executable_path=ChromeDriverManager().install(), chrome_options=chrome_options)
        driver.get('https://patents.google.com')
        xpath='//input[@id="searchInput"]'
        searchBox=driver.find_element_by_xpath(xpath)
        searchBox.send_keys(searchterm[0])
        searchBox.send_keys(Keys.ENTER)

        time.sleep(5)
        purl = driver.current_url
        print(f" Patent URL: {purl}\n")  
        response=get(purl)      
        data=response.text
        soup=BeautifulSoup(data, 'html.parser')
        driver.quit()
        
        #Claimset = form.save(commit=False)
        allclaims=soup.findAll(id=True, class_="claim")
        cln=len(allclaims) # Total claims
        print("Length of all claims is", cln)
        claimlist=[]  # to hold array of stings of claims, an empty list is declared

        n=0
        print("Before appending, num of claimset objects were", len(Claimset.objects.all()))
        while n < cln :
            claimlist.append(allclaims[n].text) # to genrate a list of strings
            Claimset.objects.create(patno=claims.patno, pubno=claims.pubno,claim_list=allclaims[n].text, total_claim=n+1, add112=AB112, obj=claims.obj)
            #Claimset.objects.update_or_create( claim_list=claims[n].text, total_claim=cln)
            n +=1
        print("After appending, num of claimset objects were", len(Claimset.objects.all()))
        print("claimlist length:", len(claimlist), claimlist)
        patnopath = os.path.join(BASE_DIR,'patoa', 'static','media')
        print("path for docx file", patnopath)
        print("path for docx file", searchterm[0])
        claims2doc(claimlist, patnopath, AB112, searchterm[0])
        #form =ClaimForm()
        return redirect('/')

        # form =ClaimForm()    
        # cln =0; 
        # cln = form.cleaned_data.get('total_claim')
        # context ={'form':form,'totalclaim':cln}
        # claimset=Claimset.objects.all()
        
        # searchterm = form.cleaned_data.get("patno")
    return render(request, 'patoa/result.html', {'form':form, 'searchterm':searchterm[0], 'data': Claimset.objects.all()[1:]})#[] if not isinstance(claims, list) else claims})
        #'searchterm':searchterm,'totalclaim':cln })
    

def clear(request):
    Claimset.objects.all().delete()
    return render(request,'patoa/result.html')

def claims2doc(claimslist, path, sel, searchterm):
    docu = Document()
    docu.styles['Normal'].font.size = Pt(12)
    docu.styles['Normal'].font.name = 'Arial'
    docu.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)

    def make_paragraph(text, size=12):
        p = docu.add_paragraph()
        p.add_run(text).font.size = Pt(size)
        #p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if sel == '112obj' :
        print(sel)
        docu.add_heading("Claim rejection under 35 USC 112")
        make_paragraph('The following is a quotation of the first paragraph of 35 U.S.C. 112(a):')
        make_paragraph('(a) IN GENERAL.—The specification shall contain a written description of the invention, and of the manner ' + \
                        'and process of making and using it, in such full, clear, concise, and exact terms as to enable any person skilled in the art to which it pertains,' + \
                        'or with which it is most nearly connected, to make and use the same, and shall set forth the best mode contemplated by the inventor or joint inventor of carrying out the invention. \n \n', size=10)
        make_paragraph('The following is a quotation of the first paragraph of pre-AIA 35 U.S.C. 112: \n')
        make_paragraph('The specification shall contain a written description of the invention, and of the manner and process of making ' + \
                'and using it, in such full, clear, concise, and exact terms as to enable any person skilled in the art to which it pertains, ' + \
                'or with which it is most nearly connected, to make and use the same, and shall set forth the best mode contemplated ' +  \
                    'by the inventor of carrying out his invention.\n', size=10)
    elif sel == '112':
        print(sel)
    elif sel == 'obj' :
        print(sel)   
    else:
        print("None")
        AB112 =''

    p=docu.add_paragraph()
    for i in range(len(claimslist)):
        claim = claimslist[i].strip()
        for j in range(len(claim)): 
            if claim[j] == ".":
                claim = claim[j:]
                break
        p.add_run(f"Regarding claim {i+1}").bold = True
        p.add_run(f"{claim} \n")
        # #print(p.runs[0].text)
        # #p.runs[0].bold = True
        # p.runs[0].font.name = 'Arial'
        #p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    all_paras = docu.paragraphs
    #print(f"Added all claims. No of Paras:", len(all_paras))
    print("doc files name prefix", searchterm)
    filename = ("/" + str(searchterm) + '.docx')
    docu.save(path+filename)

